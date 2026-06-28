from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "דוח_פרויקט_DeadZone_מעודכן.docx"
ASSETS = ROOT / "report_assets"
ASSETS.mkdir(exist_ok=True)
LOGO = ROOT / "deadzone-client" / "public" / "deadZone_Logo.png"

NAVY = "17324D"
BLUE = "285E8E"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = "9B1C1C"
WHITE = "FFFFFF"
INK = "17212B"
MUTED = "5B6670"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B8C1CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def set_rtl(paragraph, rtl=True):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1" if rtl else "0")


def set_run_font(run, name="Arial", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def set_picture_alt(inline_shape, alt_text):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def add_picture(doc, path, width, alt_text):
    shape = doc.add_picture(str(path), width=width)
    set_picture_alt(shape, alt_text)
    return shape


def add_text(doc, text, bold=False, italic=False, size=11, color=INK, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=6, rtl=True, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep
    set_rtl(p, rtl)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.45 + level * 0.35)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_callout(doc, title, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 150, 180, 150, 180)
    set_table_borders(table, color="AAB7C4", size="5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    r = p.add_run(f"{title}: ")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if row_index % 2:
                set_cell_shading(cell, "F8FAFC")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, 130, 150, 130, 150)
    set_table_borders(table, "CBD2D9", "5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_rtl(p, False)
    for idx, line in enumerate(code.strip().splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.2, color="25313C")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def diagram(title, boxes, arrows, filename, size=(1500, 820)):
    image = Image.new("RGB", size, "#F6F8FB")
    d = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    bold_path = Path("C:/Windows/Fonts/arialbd.ttf")
    font = ImageFont.truetype(str(font_path), 28)
    small = ImageFont.truetype(str(font_path), 23)
    bold = ImageFont.truetype(str(bold_path), 34)
    d.text((size[0] // 2, 45), title, fill="#17324D", font=bold, anchor="mm")
    for key, (xy, label, fill) in boxes.items():
        x1, y1, x2, y2 = xy
        d.rounded_rectangle(xy, radius=18, fill=fill, outline="#285E8E", width=4)
        lines = []
        for part in label.split("\n"):
            lines.extend(wrap(part, width=22))
        total = len(lines) * 31
        y = (y1 + y2) / 2 - total / 2 + 15
        for line in lines:
            d.text(((x1 + x2) / 2, y), line, fill="#17212B", font=small, anchor="mm")
            y += 31
    for start, end, label in arrows:
        sx, sy = start
        ex, ey = end
        d.line((sx, sy, ex, ey), fill="#4A6175", width=5)
        angle = __import__("math").atan2(ey - sy, ex - sx)
        for shift in (-0.5, 0.5):
            ax = ex - 22 * __import__("math").cos(angle + shift)
            ay = ey - 22 * __import__("math").sin(angle + shift)
            d.line((ex, ey, ax, ay), fill="#4A6175", width=5)
        if label:
            d.text(((sx + ex) / 2, (sy + ey) / 2 - 14), label, fill="#5B6670", font=font, anchor="mm")
    path = ASSETS / filename
    image.save(path, quality=95)
    return path


architecture_img = diagram(
    "DeadZone - System Architecture",
    {
        "browser": ((80, 280, 390, 500), "Browser\nReact + Three.js\nHUD and Lobby", "#E8F2FA"),
        "rest": ((560, 130, 920, 330), "REST API\nSpring Boot MVC\nUsers / Rooms / Social", "#E8EEF5"),
        "ws": ((560, 480, 920, 680), "WebSocket /game\nRealtime room state\nMOVE / SHOOT / HIT", "#FFF0E8"),
        "db": ((1090, 130, 1400, 330), "PostgreSQL / Supabase\nJPA + Flyway", "#EAF5EA"),
        "mail": ((1090, 480, 1400, 680), "SMTP Email\n6-digit verification", "#F6EDF8"),
    },
    [
        ((390, 350), (560, 230), "HTTPS + JWT"),
        ((390, 440), (560, 580), "WebSocket"),
        ((920, 230), (1090, 230), "JPA"),
        ((920, 260), (1090, 580), "SMTP"),
    ],
    "architecture.png",
)

erd_img = diagram(
    "DeadZone - ERD (Main Relations)",
    {
        "users": ((590, 300, 910, 520), "USERS\nPK id\nusername, email\nwallet, xp, admin", "#E8EEF5"),
        "tokens": ((80, 80, 390, 250), "EMAIL_VERIFICATION_TOKENS\nFK user_id", "#F6EDF8"),
        "requests": ((80, 330, 390, 500), "FRIEND_REQUESTS\nFK sender_id\nFK recipient_id", "#E8F2FA"),
        "friends": ((80, 580, 390, 750), "FRIENDSHIPS\nFK first_user_id\nFK second_user_id", "#EAF5EA"),
        "rooms": ((1110, 80, 1420, 250), "LOBBY_ROOMS\nPK id\nmap, mode, players", "#FFF0E8"),
        "invites": ((1110, 330, 1420, 500), "ROOM_INVITATIONS\nFK sender/recipient\nFK room_code", "#FFF6DD"),
        "collections": ((1020, 600, 1490, 780), "USER_* COLLECTION TABLES\ninventory, accessories,\nupgrades, missions, stats", "#EEF3F6"),
    },
    [
        ((590, 350), (390, 165), "1:N"),
        ((590, 400), (390, 415), "1:N x2"),
        ((590, 470), (390, 665), "1:N x2"),
        ((910, 395), (1110, 415), "1:N x2"),
        ((1265, 250), (1265, 330), "1:N"),
        ((910, 500), (1100, 650), "1:N"),
    ],
    "erd.png",
)

flow_img = diagram(
    "DeadZone - Screen Flow",
    {
        "loading": ((40, 300, 230, 450), "Loading", "#EEF3F6"),
        "auth": ((290, 300, 500, 450), "Authentication\nLogin / Register /\nVerify Email", "#F6EDF8"),
        "menu": ((570, 300, 800, 450), "Main Menu\nDaily Missions", "#E8EEF5"),
        "player": ((900, 70, 1140, 220), "My Player\nStore / Loadout", "#EAF5EA"),
        "settings": ((900, 270, 1140, 420), "Settings\nKey Bindings", "#EEF3F6"),
        "rooms": ((900, 470, 1140, 620), "Rooms\nSocial / Create", "#FFF6DD"),
        "match": ((1240, 470, 1460, 620), "3D Match\nHUD / Scoreboard", "#FFF0E8"),
    },
    [
        ((230, 375), (290, 375), ""),
        ((500, 375), (570, 375), "JWT"),
        ((800, 350), (900, 145), ""),
        ((800, 375), (900, 345), ""),
        ((800, 410), (900, 545), ""),
        ((1140, 545), (1240, 545), "Join"),
        ((1350, 470), (1080, 420), "Exit"),
    ],
    "screen_flow.png",
)

usecase_img = diagram(
    "DeadZone - Use Case Overview",
    {
        "user": ((40, 310, 250, 500), "Registered Player", "#E8EEF5"),
        "admin": ((40, 580, 250, 740), "Admin Account", "#F6EDF8"),
        "auth": ((430, 60, 760, 210), "Register / Verify / Login", "#EEF3F6"),
        "room": ((430, 250, 760, 400), "Find / Create / Join Room", "#FFF6DD"),
        "game": ((430, 440, 760, 590), "Play Match / Objectives", "#FFF0E8"),
        "social": ((920, 60, 1270, 210), "Friends / Room Invites", "#E8F2FA"),
        "progress": ((920, 250, 1270, 400), "XP / Missions / Rewards", "#EAF5EA"),
        "store": ((920, 440, 1270, 590), "Buy / Equip / Upgrade", "#F6EDF8"),
        "benefit": ((920, 640, 1270, 780), "Admin Benefits\nFull wallet and XP", "#FFF6DD"),
    },
    [
        ((250, 350), (430, 135), ""),
        ((250, 385), (430, 325), ""),
        ((250, 420), (430, 515), ""),
        ((250, 350), (920, 135), ""),
        ((250, 400), (920, 325), ""),
        ((250, 450), (920, 515), ""),
        ((250, 650), (920, 710), ""),
    ],
    "use_cases.png",
)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for level, size, color, before, after in [
    (1, 16, BLUE, 16, 8),
    (2, 13, BLUE, 12, 6),
    (3, 11.5, NAVY, 8, 4),
]:
    s = styles[f"Heading {level}"]
    s.font.name = "Arial"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    s._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_rtl(header)
run = header.add_run("DeadZone | דוח סיכום ותכנון פרויקט")
set_run_font(run, size=9, color=MUTED)
add_page_number(section.footer.paragraphs[0])

# Cover
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(6.4))
    set_picture_alt(logo_shape, "לוגו מערכת DeadZone")
    p.paragraph_format.space_after = Pt(18)
add_text(doc, "דוח סיכום ותכנון פרויקט תוכנה", bold=True, size=25, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
add_text(doc, "DeadZone — מערכת משחק יריות תלת־ממדית מרובת משתתפים", bold=True, size=17,
         color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
add_table(doc, ["פרט", "ערך"], [
    ["שם הפרויקט", "DeadZone"],
    ["שם המערכת", "DeadZone Multiplayer 3D Shooter"],
    ["שם הקורס", "לא זוהה בקוד — יש להשלים ידנית"],
    ["שמות מגישים", "לא זוהו בקוד — יש להשלים ידנית"],
    ["תאריך הגשה", "25 ביוני 2026"],
    ["גרסת הדוח", "1.0 — מבוססת על סריקת קוד הפרויקט"],
], widths=[1.7, 4.65], font_size=10)
add_callout(doc, "הערת מתודולוגיה", "הדוח נכתב על סמך קוד המקור, קובצי התצורה, סכמת מסד הנתונים, בדיקות היחידה ונכסי הממשק שנמצאו בפרויקט. מידע שלא ניתן היה לאמת סומן להשלמה ידנית.")
doc.add_page_break()

# TOC
add_heading(doc, "תוכן עניינים", 1)
toc = [
    "תקציר הדוח", "פרק 1 — שם ותיאור כללי של המערכת", "פרק 2 — מערכות דומות קיימות",
    "פרק 3 — דרישות פונקציונליות", "פרק 4 — דרישות לא פונקציונליות",
    "פרק 5 — ישויות וקשרים", "פרק 6 — תסריטי שימוש (Use Cases)",
    "פרק 7 — תהליכים עסקיים וטכניים", "פרק 8 — מסד נתונים",
    "פרק 9 — עיצוב תפריטים", "פרק 10 — עיצוב ויזואלי של מסכים",
    "פרק 11 — תרשים זרימה בין מסכים", "פרק 12 — שאילתות ודוחות",
    "פרק 13 — בדיקות QA", "פרק 14 — טכנולוגיות",
    "פרק 15 — קריטריונים להצלחת המערכת", "פרק 16 — לוח זמנים",
    "פרק 17 — נספחים, תרשימים וביבליוגרפיה", "פרק 18 — English Abstract",
    "רשימת פרטים להשלמה ידנית",
]
add_numbered(doc, toc)
doc.add_page_break()

add_heading(doc, "תקציר הדוח", 1)
add_text(doc, "DeadZone היא מערכת משחק יריות בגוף ראשון הפועלת בדפדפן ומחברת בין ממשק תלת־ממדי אינטראקטיבי, שירותי משתמשים וחברה, חדרי משחק ותקשורת בזמן אמת. המערכת מיועדת לשחקנים המעוניינים להיכנס במהירות למשחק קצר, לבחור מפה ומצב משחק, להתקדם ברמות, להשלים משימות, לרכוש פריטים קוסמטיים ולשחק עם חברים.")
add_text(doc, "הבעיה המרכזית שהמערכת פותרת היא שילוב חוויית FPS עשירה ונגישה ללא התקנת לקוח ייעודי, תוך שמירת זהות והתקדמות שחקן בצד השרת. המימוש מפריד בין צד לקוח ב־React ו־Three.js, צד שרת ב־Spring Boot, מסד PostgreSQL מנוהל באמצעות JPA ו־Flyway, ומנגנון WebSocket לעדכוני מצב משחק.")
add_text(doc, "יתרונה המרכזי של DeadZone הוא רוחב המערכת ביחס לפרויקט דפדפן: חמישה מצבי משחק, תשע מפות, בוטים, נשקים וסקינים, משימות יומיות, כלכלה פנימית, חנות, התאמת דמות, מערכת חברים והזמנות, אימות דוא״ל ו־JWT — כולם מחוברים למודל התקדמות מתמשך.")
add_table(doc, ["היבט", "סיכום"], [
    ["קהל יעד", "שחקני דפדפן, תלמידים וחובבי משחקי FPS קצרים מרובי־משתתפים"],
    ["צד לקוח", "React 19, Vite 8, Three.js, React Three Fiber/Drei/Rapier"],
    ["צד שרת", "Java 21, Spring Boot 4.1, Spring MVC, WebSocket"],
    ["מסד נתונים", "PostgreSQL/Supabase בפריסה; H2 במצב offline ובבדיקות"],
    ["אבטחה", "JWT מסוג HS256, סיסמאות SHA-256 עם salt, אימות דוא״ל בן 6 ספרות"],
    ["יכולות בולטות", "חדרים, חברים, הזמנות, חמישה מצבי משחק, משימות, XP, חנות והתאמה אישית"],
], widths=[1.55, 4.8], font_size=9.5)

add_heading(doc, "פרק 1 — שם ותיאור כללי של המערכת", 1)
add_heading(doc, "1.1 שם המערכת", 2)
add_text(doc, "שם המערכת הוא DeadZone. השם מופיע בנכסי המיתוג, בשמות חבילות צד הלקוח ובמבנה חבילת השרת.")
add_heading(doc, "1.2 תיאור כללי", 2)
add_text(doc, "המערכת היא יישום SPA בדפדפן הכולל לובי ומשחק תלת־ממדי. השחקן יוצר חשבון, מאמת כתובת דוא״ל, מתחבר, בוחר ציוד ודמות, נכנס לחדר קיים או יוצר חדר חדש, ומשחק מול שחקנים אחרים או בוטים. נתוני החשבון וההתקדמות נשמרים במסד הנתונים; מצב הקרב הפעיל מסונכרן באמצעות WebSocket.")
add_heading(doc, "1.3 מטרות", 2)
add_bullets(doc, [
    "לאפשר כניסה מהירה למשחק FPS מתוך דפדפן מודרני.",
    "לספק חוויית משחק חוזרת באמצעות רמות, XP, מטבע, משימות יומיות ופתיחת תוכן.",
    "לתמוך במשחק חברתי באמצעות חיפוש שחקנים, בקשות חברות והזמנות לחדר.",
    "להציע מגוון מצבי משחק: Team Deathmatch, Free For All, Capture the Flag, Attack/Defend ו־Circle Control.",
    "לשמור מידע רגיש והתקדמות בצד השרת ולבצע בדיקות תקינות על רכישות ועדכוני נתונים.",
])
add_heading(doc, "1.4 קהל יעד ותהליכים משופרים", 2)
add_text(doc, "קהל היעד הוא משתמשים בעלי דפדפן התומך ב־WebGL ובקלט מקלדת/עכבר. המערכת מרכזת תחת ממשק יחיד תהליכים שבמשחקים פשוטים מתבצעים ללא התמדה: יצירת זהות, התאמת מקשים, ניהול ציוד, יצירת חדר פרטי, בחירת מצב משחק, הזמנת חברים, מעקב משימות וקבלת תגמולים.")

add_heading(doc, "פרק 2 — מערכות דומות קיימות", 1)
add_text(doc, "ההשוואה בפרק זה היא פונקציונלית בלבד. מאפייני DeadZone מבוססים על הקוד; מאפייני המערכות החיצוניות מתוארים ברמה כללית לפי אתרי המוצר הרשמיים המופיעים בביבליוגרפיה.")
add_table(doc, ["מערכת", "מאפיינים דומים", "הבדלים ביחס ל־DeadZone"], [
    ["Counter-Strike 2", "משחק יריות קבוצתי, מפות, נשקים ומטרות תחרותיות.", "מערכת מסחרית רחבת היקף המותקנת במחשב; DeadZone ממוקדת בדפדפן ובחדרים של עד 6 שחקנים."],
    ["VALORANT", "קרבות קבוצתיים, מפות ייעודיות, מצבי Attack/Defend ופרופיל שחקן.", "VALORANT מבוססת סוכנים ויכולות ייחודיות; DeadZone מבוססת נשקים, power-ups ומצבי משחק קלים יותר."],
    ["Krunker", "FPS בדפדפן, משחק מהיר, ציוד, קוסמטיקה ומצבי משחק.", "Krunker היא מערכת מסחרית ותיקה; DeadZone מציגה מימוש לימודי מלא ושקוף עם Spring Boot, JPA ו־WebSocket."],
    ["Fortnite", "קוסמטיקה, XP, משימות, מטבע פנימי, בוטים ומגוון מצבים.", "Fortnite פועלת בקנה מידה ומודל Battle Royale/פלטפורמה; DeadZone היא זירת FPS קטנה וממוקדת."],
], widths=[1.25, 2.35, 2.75], font_size=8.6)
add_heading(doc, "2.1 ייחוד המערכת", 2)
add_text(doc, "הייחוד של DeadZone אינו מנגנון משחק יחיד, אלא חיבור של שכבות רבות בפרויקט אחד: מנוע תלת־ממדי בדפדפן, מצבי משחק אובייקטיביים, שירות משתמשים מלא, כלכלה מאומתת בצד השרת, מערכת חברתית ומודל נתונים מנורמל. הקוד כולל גם מצב offline מבוסס H2, המקל על פיתוח והדגמה ללא שירות ענן.")

add_heading(doc, "פרק 3 — דרישות פונקציונליות", 1)
functional = [
    ("RF-01", "רישום משתמש", "קליטת שם משתמש, דוא״ל וסיסמה; יצירת משתמש והפקת קוד אימות."),
    ("RF-02", "אימות דוא״ל", "שליחת קוד בן 6 ספרות, תוקף 15 דקות וסימון הקוד כמשומש."),
    ("RF-03", "התחברות", "בדיקת סיסמה, חסימת חשבון לא מאומת והנפקת JWT ל־30 יום."),
    ("RF-04", "שחזור סשן", "טעינת המשתמש הנוכחי באמצעות GET /api/users/me."),
    ("RF-05", "שמירת התקדמות", "שמירת ארנק, XP, K/A/D, ציוד, שדרוגים ונתוני משימות."),
    ("RF-06", "צפייה בחדרים", "שליפת חדרים פתוחים ורענון הלובי כל 3 שניות."),
    ("RF-07", "איתור חדר בקוד", "חיפוש חדר באמצעות קוד הזמנה."),
    ("RF-08", "יצירת חדר", "בחירת שם, מפה, מצב משחק, יעד ניקוד, זמן, קיבולת ובוטים."),
    ("RF-09", "כניסה ויציאה מחדר", "עדכון מונה שחקנים וחלוקת קבוצות אוטומטית."),
    ("RF-10", "משחק בזמן אמת", "JOIN, MOVE, SHOOT, HIT והפצת STATE/SHOT דרך WebSocket."),
    ("RF-11", "מצבי משחק", "TDM, FFA, CTF, Attack/Defend ו־Circle Control."),
    ("RF-12", "מנוע קרב", "תנועה, קפיצה, ירי, פגיעה, רתע, טעינה, רימונים, מוות וחזרה למשחק."),
    ("RF-13", "בוטים", "מילוי מקומות פנויים בבוטים כאשר האפשרות מופעלת."),
    ("RF-14", "Power-ups", "בריאות, נזק, מגן, מהירות וקצב אש."),
    ("RF-15", "חנות", "רכישה וציוד של תלבושות, סקינים, אביזרים ושדרוגי נשק."),
    ("RF-16", "פתיחת תוכן", "נעילת נשקים ומפות לפי רמת השחקן."),
    ("RF-17", "משימות יומיות", "שלוש משימות פעילות, מעקב התקדמות, קבלת פרס ו־reroll."),
    ("RF-18", "מערכת חברתית", "חיפוש שחקנים, שליחת/קבלת/דחיית בקשת חברות."),
    ("RF-19", "הזמנות לחדר", "הזמנת חבר לחדר וקבלת או דחיית הזמנה עד 30 דקות."),
    ("RF-20", "התאמת מקשים", "שינוי מקשי השליטה ושמירתם ב־localStorage לפי משתמש."),
    ("RF-21", "HUD ולוח תוצאות", "הצגת חיים, תחמושת, רימונים, כסף, זמן, ניקוד ו־K/A/D."),
    ("RF-22", "תפריטי עצירה ומוות", "עצירה, התאמת דמות, החלפת נשק ויציאה מאושרת."),
    ("RF-23", "חשבון מנהל", "משתמש admin מקבל ארנק/XP גבוהים וכלל פריטי החנות."),
]
add_table(doc, ["מזהה", "דרישה", "פירוט"], functional, widths=[0.75, 1.65, 3.95], font_size=8.5)
add_callout(doc, "הבהרה", "לא נמצא ממשק ניהול משתמשים ייעודי המאפשר למנהל ליצור, לערוך או למחוק משתמשים. קיים Role מסוג admin, אך השפעתו בקוד היא בעיקר הטבות במשחק.")

add_heading(doc, "פרק 4 — דרישות לא פונקציונליות", 1)
add_table(doc, ["תחום", "דרישה מומלצת ומצב שנמצא בקוד"], [
    ["ביצועים", "רינדור WebGL בצד הלקוח; עדכון חדרים כל 3 שניות; הגבלת pixel ratio ל־2. יעד מומלץ: תגובת REST עד 500ms ו־FPS יציב של 45–60 במחשב יעד."],
    ["אבטחת מידע", "JWT חתום HS256; סיסמה עם salt ו־SHA-256; אימות דוא״ל; בדיקות מחיר וגבולות התקדמות בצד שרת. חסר: Spring Security, rate limiting והגנת Origin מצומצמת."],
    ["שימושיות", "מסכי טעינה, לובי, HUD, לוח תוצאות, תפריט עצירה והודעות מצב. הממשק בפועל באנגלית; עברית לא נדרשה בקוד."],
    ["זמינות", "מצב production עם PostgreSQL ומצב offline עם H2. לא זוהו health checks, ניטור או SLA."],
    ["תחזוקתיות", "הפרדה בין controllers, services, repositories, models ומודולי משחק בצד הלקוח. עם זאת, useDeadzoneController ו־Lobby.jsx גדולים וצפופים."],
    ["תאימות", "יישום דפדפן מבוסס WebGL. CSS כולל התאמות responsive, אך השליטה מבוססת מקלדת ועכבר ולכן תמיכה מלאה במובייל לא זוהתה."],
    ["מדרגיות", "REST ומסד נתונים יכולים להתרחב אופקית, אך חדרי WebSocket נשמרים בזיכרון תהליך יחיד ולכן אינם משותפים בין מופעים."],
    ["גיבוי ושחזור", "ב־PostgreSQL/Supabase ניתן לבצע גיבוי ברמת הפלטפורמה, אך לא נמצאה מדיניות גיבוי או תהליך שחזור בקוד."],
    ["נגישות", "קיימים aria-label בחלק מהפקדים ודיאלוגים, אך לא נמצאה בדיקת נגישות מלאה או ניווט מקלדת לכל המסכים."],
    ["אמינות נתונים", "Flyway, מפתחות זרים, unique indexes ו־optimistic locking באמצעות @Version. מצב offline משתמש ddl-auto=update."],
], widths=[1.35, 5.0], font_size=8.8)
add_heading(doc, "4.1 ביצועים וקצב עדכון — ממצאים מהקוד", 2)
add_table(doc, ["רכיב", "מה נמצא בקוד", "מסקנה"], [
    ["Game Loop", "GameWorld.animate מופעל באמצעות requestAnimationFrame וקורא לעדכון, סנכרון meshes ורינדור בכל מסגרת.", "קיים Game Loop מבוסס קצב הרענון של הדפדפן."],
    ["FPS Target", "לא נמצא ערך יעד FPS, מונה FPS או מנגנון נעילת קצב מסגרות.", "לא נמצא בקוד; יש לבצע מדידות בפועל."],
    ["Render Loop", "THREE.WebGLRenderer.render נקרא בכל requestAnimationFrame. גם CharacterPreview משתמש בלולאת requestAnimationFrame.", "הרינדור רציף ותלוי בדפדפן ובחומרה."],
    ["Delta Time", "הפרש הזמן נלקח מ־THREE.Clock ומוגבל ל־0.033 שניות לכל צעד.", "ההגבלה מונעת קפיצת סימולציה גדולה; היא אינה הוכחה ליעד של 30 FPS."],
    ["Physics Tick", "תנועה, בוטים, התנגשות, קליעים, רימונים ומטרות מתעדכנים מתוך לולאת הרינדור באמצעות dt.", "אין tick פיזיקה קבוע ונפרד במנוע המשחק המרכזי."],
    ["Rapier", "RigidBody קבוע מופיע ברכיבי מפת ApocalypticArena, אך GameWorld משתמש בעיקר במערכות CollisionSystem ו־PlayerCollisionSystem מותאמות.", "קיימת תלות במנוע Rapier, אך הפיזיקה המרכזית אינה מנוהלת כולה על ידו."],
    ["WebSocket Movement", "RealtimeClient מגדיר SEND_INTERVAL_MS=50 ושולח MOVE לכל היותר אחת ל־50ms.", "קצב מרבי מתועד: 20 עדכוני תנועה בשנייה."],
    ["אירועי פגיעה", "HIT נשלח מיידית ללא throttle; השרת מפיץ מצב חדר לאחר טיפול בהודעה.", "קצב תלוי באירועי המשחק, ללא tick שרת קבוע."],
    ["Lobby Update", "רענון חדרים ומידע חברתי מתבצע כל 3,000ms.", "קצב polling מתועד: אחת ל־3 שניות."],
    ["Latency Handling", "לא נמצאו ping/pong יישומי, מדידת RTT, interpolation, buffer snapshots, reconnect/backoff או lag compensation.", "לא קיים בפרויקט הנוכחי; מומלץ להוספה בעתיד."],
], widths=[1.25, 3.45, 1.65], font_size=8.0)
add_callout(doc, "מדידות ביצועים", "הקוד אינו כולל תוצאות FPS, latency, עומס CPU/GPU או זמן תגובה של API. יש לבצע מדידות בפועל בסביבת היעד.")
add_heading(doc, "4.2 אבטחת Production — מצב מאומת", 2)
add_table(doc, ["נושא", "מצב בפרויקט", "פירוט מבוסס קוד"], [
    ["bcrypt / Argon2", "לא קיים", "PasswordService משתמש ב־SHA-256 עם salt אקראי של 16 בתים. מומלץ להחליף ב־bcrypt או Argon2."],
    ["JWT", "קיים", "JwtService מותאם יוצר JWT חתום HS256 עם sub, username ו־exp ל־30 יום."],
    ["CORS", "קיים באופן פתוח", "@CrossOrigin(origins=\"*\") ב־Controllers ו־setAllowedOrigins(\"*\") ב־WebSocket."],
    ["CSRF", "לא נמצא בקוד", "אין Spring Security ואין תצורת CSRF."],
    ["HTTPS", "לא נאכף בשרת", "api/config ממיר https ל־wss כאשר ה־origin כבר HTTPS; אין redirect או certificate configuration בקוד."],
    ["Validation", "קיימת ידנית וחלקית", "Services בודקים שדות, טווחים, מחירים וקלט; לא נמצאו @Valid או Bean Validation annotations."],
    ["Spring Security", "לא קיים", "לא קיימת dependency או SecurityFilterChain."],
    ["Authentication", "קיים", "Registration, email verification, login, password matching ו־JWT Bearer."],
    ["Authorization", "חלקי", "כל endpoint מוגן מפיק userId מה־JWT; אין מערכת permissions ואין endpoints המוגנים לפי Role."],
    ["Admin", "קיים כדגל נתונים", "User.admin ו־seed למשתמש test; ההשפעה היא הטבות XP/ארנק/מלאי, לא ממשק ניהול."],
    ["Rate Limiting", "לא קיים", "לא נמצאה הגבלת בקשות לרישום, login, אימות, REST או WebSocket."],
    ["WebSocket Security", "לא קיים", "ה־handshake אינו דורש JWT, ה־origin פתוח, וזהות playerId נשלחת מהלקוח."],
], widths=[1.25, 1.25, 3.85], font_size=8.0)

add_heading(doc, "פרק 5 — ישויות וקשרים", 1)
add_text(doc, "המערכת כוללת ישויות מתמידות במסד הנתונים וישויות זמן־ריצה המשמשות את מנוע המשחק. הטבלה הבאה מרכזת את הישויות העיקריות.")
add_table(doc, ["ישות", "תיאור", "שדות מרכזיים", "קשרים"], [
    ["User", "חשבון שחקן, התקדמות, ציוד והרשאה.", "id, username, email, password, wallet, xp, admin", "1:N לכל טבלאות האוסף; 1:N לטוקנים, בקשות, חברויות והזמנות."],
    ["EmailVerificationToken", "קוד אימות דוא״ל זמני.", "id, token, user_id, expires_at, used_at", "N:1 User."],
    ["LobbyRoom", "חדר זמין בלובי.", "id, name, map_id, game_mode, players, max_players", "1:N RoomInvitation."],
    ["FriendRequest", "בקשת חברות מכוונת.", "sender_id, recipient_id, status, created_at", "שני FK אל User."],
    ["Friendship", "קשר חברות בלתי־מכוון המאוחסן בסדר מזהים.", "first_user_id, second_user_id, created_at", "שני FK אל User."],
    ["RoomInvitation", "הזמנה מחבר לחדר.", "sender_id, recipient_id, room_code, status, expires_at", "N:1 לשני Users ול־LobbyRoom."],
    ["GameRoom", "חדר בזמן אמת בזיכרון השרת.", "id, mapId, playersBySessionId", "1:N Player; אינו Entity של JPA."],
    ["Player", "מצב שחקן פעיל במשחק.", "position, team, health, kills, assists, deaths, score", "שייך ל־GameRoom בזמן ריצה."],
    ["GameMessage", "הודעת WebSocket.", "type, roomId, playerId, transform, health, score", "DTO זמן־ריצה ללא טבלה."],
], widths=[1.25, 1.8, 2.15, 1.2], font_size=8.1)
add_picture(doc, erd_img, Inches(6.35), "תרשים קשרי ישויות מרכזיים של DeadZone")
add_text(doc, "תרשים 1 — ERD מרכזי של DeadZone", bold=True, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
add_code(doc, """
erDiagram
    USERS ||--o{ EMAIL_VERIFICATION_TOKENS : verifies
    USERS ||--o{ FRIEND_REQUESTS : sends
    USERS ||--o{ FRIEND_REQUESTS : receives
    USERS ||--o{ FRIENDSHIPS : participates
    USERS ||--o{ ROOM_INVITATIONS : sends_or_receives
    LOBBY_ROOMS ||--o{ ROOM_INVITATIONS : targets
    USERS ||--o{ USER_OWNED_OUTFITS : owns
    USERS ||--o{ USER_WEAPON_UPGRADES : upgrades
    USERS ||--o{ USER_MAP_PLAYS : records
""")

add_heading(doc, "פרק 6 — תסריטי שימוש (Use Cases)", 1)
add_picture(doc, usecase_img, Inches(6.35), "תרשים תרחישי שימוש מרכזיים של DeadZone")
add_text(doc, "תרשים 2 — סקירת תרחישי שימוש", bold=True, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
use_cases = [
    ("UC-01 — רישום ואימות", "אורח", "יצירת חשבון מאומת",
     ["בחירת Register.", "הזנת שם משתמש, דוא״ל וסיסמה.", "השרת יוצר משתמש וקוד אימות.", "המשתמש מזין את הקוד.", "השרת מסמן את הדוא״ל כמאומת."],
     "חשבון פעיל המוכן להתחברות."),
    ("UC-02 — התחברות ושחזור סשן", "משתמש רשום", "כניסה למערכת",
     ["הזנת שם משתמש וסיסמה.", "השרת משווה hash.", "נבדק אימות הדוא״ל.", "מוחזר JWT ונשמר בדפדפן.", "בטעינה הבאה נטען /me."],
     "הצגת התפריט הראשי ונתוני השחקן."),
    ("UC-03 — יצירת חדר", "שחקן", "פתיחת חדר מותאם",
     ["מעבר למסך Rooms.", "בחירת Create Room.", "בחירת מפה, מצב, זמן, יעד, קיבולת ובוטים.", "שליחת POST לשרת.", "השרת מייצר קוד ושומר חדר."],
     "חדר חדש מופיע ונבחר בלובי."),
    ("UC-04 — כניסה לחדר ומשחק", "שחקן", "התחלת משחק",
     ["בחירת חדר.", "בדיקת זמינות וקיבולת.", "POST join ועדכון קבוצות.", "פתיחת מסך המשחק.", "יצירת GameWorld וחיבור WebSocket.", "שליחת JOIN ומצב תנועה."],
     "השחקן מופיע בזירה ומקבל מצב חדר."),
    ("UC-05 — קרב ופגיעה", "שני שחקנים", "ביצוע ירי מתוזמן",
     ["השחקן יורה.", "הלקוח מחשב פגיעה ושולח SHOOT/HIT.", "השרת מוודא שאין friendly fire.", "החיים מופחתים.", "במוות מתעדכנים kills/deaths/score.", "מופץ STATE חדש."],
     "HUD ולוח התוצאות מתעדכנים."),
    ("UC-06 — השלמת משימה", "שחקן", "קבלת תגמול יומי",
     ["פעולות משחק מעדכנות נתוני משימה.", "כרטיס המשימה מסומן Ready.", "השחקן לוחץ Claim.", "מתווספים כסף ו־XP.", "המשימה מסומנת claimed ונשמרת."],
     "התגמול נשמר בפרופיל."),
    ("UC-07 — רכישה והצטיידות", "שחקן", "קניית פריט בחנות",
     ["בחירת פריט.", "בדיקת יתרת ארנק.", "הלקוח שולח מלאי וארנק מעודכנים.", "השרת מחשב מחיר לפי קטלוג.", "השרת דוחה חוסר התאמה או שומר רכישה.", "הפריט מצויד."],
     "הפריט בבעלות השחקן והארנק מופחת."),
    ("UC-08 — שדרוג נשק", "שחקן", "העלאת רמת נשק",
     ["בחירת נשק.", "חישוב מחיר לפי הרמה.", "שליחת רמת שדרוג חדשה.", "השרת מגביל עד 10 ומחשב את המחיר המצטבר.", "שמירה במסד."],
     "רמת הנשק נשמרת."),
    ("UC-09 — בקשת חברות", "שני שחקנים", "יצירת קשר חברתי",
     ["חיפוש לפי שם משתמש.", "שליחת בקשה.", "הנמען רואה בקשה נכנסת.", "קבלה יוצרת Friendship; דחייה מוחקת בקשה."],
     "רשימת החברים מתעדכנת."),
    ("UC-10 — הזמנה לחדר", "חברים", "צירוף חבר למשחק",
     ["בחירת חדר.", "לחיצה Invite ליד חבר.", "יצירת RoomInvitation בתוקף 30 דקות.", "החבר מקבל או דוחה.", "בקבלה החדר נטען ונבחר."],
     "החבר מוכן להצטרף לחדר."),
    ("UC-11 — התאמת מקשים", "שחקן", "שינוי שליטה",
     ["מעבר להגדרות.", "בחירת פעולה.", "לחיצה על מקש חדש.", "שמירה ב־localStorage תחת מפתח המשתמש.", "טעינה במשחק הבא."],
     "המשחק משתמש במיפוי החדש."),
    ("UC-12 — עצירה, התאמה וחזרה", "שחקן במשחק", "שינוי דמות תוך כדי משחק",
     ["פתיחת Pause.", "הקפאת GameWorld המקומי.", "בחירת תלבושת/נשק/אביזר בבעלות.", "חזרה למשחק.", "המצב החדש נשלח בהודעות תנועה."],
     "המשחק ממשיך עם ציוד מעודכן."),
]
for name, actors, goal, steps, result in use_cases:
    add_heading(doc, name, 2)
    add_table(doc, ["שדה", "תוכן"], [
        ["שחקנים", actors], ["מטרה", goal], ["תהליך", " → ".join(steps)], ["תוצאה צפויה", result]
    ], widths=[1.25, 5.1], font_size=8.7)

add_heading(doc, "פרק 7 — תהליכים עסקיים וטכניים", 1)
add_picture(doc, architecture_img, Inches(6.35), "תרשים ארכיטקטורת מערכת DeadZone")
add_text(doc, "תרשים 3 — ארכיטקטורת המערכת", bold=True, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading(doc, "7.1 זרימת משתמש", 2)
add_numbered(doc, [
    "טעינת נכסים ובדיקת זמינות השרת.",
    "רישום/התחברות ואימות זהות.",
    "טעינת פרופיל, משימות, מלאי, חדרים ונתונים חברתיים.",
    "בחירת מסך: שחקן, הגדרות או חדרים.",
    "יצירה או בחירה של חדר והצטרפות.",
    "הפעלת GameWorld, חיבור WebSocket וניהול המשחק.",
    "עדכון התקדמות, משימות וארנק ב־REST.",
    "יציאה מהחדר וחזרה ללובי.",
])
add_heading(doc, "7.2 תקשורת לקוח–שרת", 2)
add_text(doc, "פעולות חשבון, חדרים וחברה מתבצעות באמצעות JSON מעל HTTP. לאחר התחברות, הלקוח מוסיף Authorization: Bearer <token>. המשחק הפעיל משתמש בחיבור WebSocket בנתיב /game. הלקוח שולח JOIN, MOVE, SHOOT ו־HIT; השרת מפיץ STATE ו־SHOT או ERROR.")
add_heading(doc, "7.3 שמירה ושליפה", 2)
add_text(doc, "Repositories של Spring Data JPA מבצעים גישה ל־PostgreSQL. Flyway יוצר את הסכמה ומוסיף מפתחות זרים ואינדקסים. אוספים של User ממופים באמצעות @ElementCollection לטבלאות נפרדות. מצב חדר לובי נשמר במסד, בעוד GameRoom ו־Player של משחק פעיל מוחזקים בזיכרון ConcurrentHashMap.")
add_heading(doc, "7.4 אימות והרשאות", 2)
add_text(doc, "בעת התחברות השרת מנפיק JWT חתום HMAC-SHA256. JwtService בודק מבנה, חתימה ותאריך תפוגה, ומחזיר userId. לא נעשה שימוש ב־Spring Security Filter Chain; כל Controller קורא ידנית ל־requireUserId. משתמש admin קיים כשדה ב־User, אולם לא נמצאו endpoints מנהליים נפרדים.")
add_heading(doc, "7.5 בקרות עסקיות", 2)
add_bullets(doc, [
    "השרת מחשב את עלות הפריטים החדשים לפי StoreCatalog ודוחה מניפולציה בארנק.",
    "עליית ארנק ללא רכישה מוגבלת ל־500 בעדכון; XP מוגבל לתוספת של 1,000.",
    "Kills, assists ו־deaths מוגבלים לעליות סבירות בעדכון יחיד.",
    "שדרוג נשק מוגבל לרמה 10.",
    "חדר מוגבל ל־2–6 שחקנים; חדר זמני ריק נמחק לאחר 5 דקות.",
    "שמות שחקנים בזמן אמת נחתכים ל־16 תווים; קבוצה מנורמלת ל־red/blue.",
])
add_code(doc, """
flowchart LR
    UI[React UI] -->|REST + JWT| API[Spring Controllers]
    UI -->|WebSocket| WS[GameWebSocketHandler]
    API --> SVC[Services]
    SVC --> JPA[Spring Data JPA]
    JPA --> DB[(PostgreSQL)]
    API --> MAIL[SMTP verification]
    WS --> MEM[In-memory GameRoom state]
""")

add_heading(doc, "פרק 8 — מסד נתונים", 1)
add_text(doc, "סכמת הייצור מוגדרת ב־V1__deadzone_schema.sql. PostgreSQL הוא מסד היעד; H2 משמש מצב offline ובדיקות. כל הטבלאות להלן זוהו בקוד.")
schema = {
    "users": [
        ("id", "BIGINT IDENTITY", "PK", "מזהה משתמש"),
        ("username", "VARCHAR(32)", "NOT NULL, UNIQUE INDEX", "שם כניסה"),
        ("email", "VARCHAR(120)", "NOT NULL", "כתובת דוא״ל; אינה Unique בקוד"),
        ("password", "VARCHAR(255)", "NOT NULL", "hash סיסמה עם salt"),
        ("email_verified", "BOOLEAN", "NOT NULL, DEFAULT false", "מצב אימות"),
        ("email_verified_at", "TIMESTAMPTZ", "NULL", "מועד אימות"),
        ("total_kills/assists/deaths", "INTEGER", "NOT NULL, DEFAULT 0", "סטטיסטיקות מצטברות"),
        ("wallet", "INTEGER", "NOT NULL, DEFAULT 0", "מטבע פנימי"),
        ("xp", "INTEGER", "NOT NULL, DEFAULT 0", "נקודות ניסיון"),
        ("outfit_id/weapon_id", "VARCHAR(255)", "NULL", "ציוד נבחר"),
        ("weapon_skin_id/grenade_skin_id", "VARCHAR(255)", "NULL", "סקינים נבחרים"),
        ("mission_stats_json", "VARCHAR(8000)", "NULL", "מצב משימות מלא ב־JSON"),
        ("admin", "BOOLEAN", "NOT NULL, DEFAULT false", "דגל מנהל"),
        ("version", "BIGINT", "NOT NULL, DEFAULT 0", "Optimistic locking"),
    ],
    "email_verification_tokens": [
        ("id", "BIGINT IDENTITY", "PK", "מזהה קוד"),
        ("token", "VARCHAR(96)", "NOT NULL, UNIQUE", "קוד אימות"),
        ("user_id", "BIGINT", "FK users(id), NOT NULL, CASCADE", "בעל הקוד"),
        ("expires_at", "TIMESTAMPTZ", "NOT NULL", "מועד תפוגה"),
        ("used_at", "TIMESTAMPTZ", "NULL", "מועד שימוש"),
    ],
    "lobby_rooms": [
        ("id", "VARCHAR(16)", "PK", "קוד חדר"),
        ("name", "VARCHAR(255)", "NOT NULL", "שם חדר"),
        ("map_id", "VARCHAR(32)", "NOT NULL", "מפה"),
        ("game_mode", "VARCHAR(32)", "NULL", "מצב משחק"),
        ("score_limit", "INTEGER", "NOT NULL, DEFAULT 30", "יעד ניקוד"),
        ("time_limit_minutes", "INTEGER", "NOT NULL, DEFAULT 20", "זמן משחק"),
        ("players/max_players", "INTEGER", "NOT NULL", "תפוסה וקיבולת"),
        ("blue_players/red_players", "INTEGER", "NOT NULL", "חלוקת קבוצות"),
        ("allow_bots/permanent", "BOOLEAN", "NOT NULL", "בוטים וחדר קבוע"),
        ("last_activity_at", "TIMESTAMPTZ", "NULL", "פעילות אחרונה"),
        ("version", "BIGINT", "NOT NULL, DEFAULT 0", "Optimistic locking"),
    ],
    "friend_requests": [
        ("id", "BIGINT IDENTITY", "PK", "מזהה בקשה"),
        ("sender_id", "BIGINT", "FK users(id), NOT NULL", "שולח"),
        ("recipient_id", "BIGINT", "FK users(id), NOT NULL", "נמען"),
        ("status", "VARCHAR(16)", "NOT NULL", "PENDING וכדומה"),
        ("created_at", "TIMESTAMPTZ", "NOT NULL", "מועד יצירה"),
        ("sender_id + recipient_id", "—", "UNIQUE", "מניעת בקשה כפולה באותו כיוון"),
    ],
    "friendships": [
        ("id", "BIGINT IDENTITY", "PK", "מזהה חברות"),
        ("first_user_id", "BIGINT", "FK users(id), NOT NULL", "המזהה הנמוך"),
        ("second_user_id", "BIGINT", "FK users(id), NOT NULL", "המזהה הגבוה"),
        ("created_at", "TIMESTAMPTZ", "NOT NULL", "מועד יצירה"),
        ("first_user_id + second_user_id", "—", "UNIQUE", "מניעת חברות כפולה"),
    ],
    "room_invitations": [
        ("id", "BIGINT IDENTITY", "PK", "מזהה הזמנה"),
        ("sender_id", "BIGINT", "FK users(id), NOT NULL", "שולח"),
        ("recipient_id", "BIGINT", "FK users(id), NOT NULL", "נמען"),
        ("room_code", "VARCHAR(16)", "FK lobby_rooms(id), NOT NULL", "חדר יעד"),
        ("status", "VARCHAR(16)", "NOT NULL", "מצב הזמנה"),
        ("created_at", "TIMESTAMPTZ", "NOT NULL", "מועד יצירה"),
        ("expires_at", "TIMESTAMPTZ", "NOT NULL, DEFAULT +30 min", "תפוגה"),
    ],
}
collection_tables = [
    ("user_owned_outfits", "sort_order", "outfit_id", "תלבושות בבעלות"),
    ("user_owned_weapon_skins", "sort_order", "skin_id", "סקיני נשק בבעלות"),
    ("user_owned_grenade_skins", "sort_order", "skin_id", "סקיני רימון בבעלות"),
    ("user_owned_accessories", "sort_order", "accessory_id", "אביזרים בבעלות"),
    ("user_equipped_accessories", "sort_order", "accessory_id", "אביזרים מצוידים"),
    ("user_claimed_missions", "sort_order", "mission_id", "משימות שנתבעו"),
]
for table_name, order_col, item_col, desc in collection_tables:
    schema[table_name] = [
        ("user_id", "BIGINT", "PK(part), FK users(id), CASCADE", "בעל הרשומה"),
        (order_col, "INTEGER", "PK(part), NOT NULL", "סדר באוסף"),
        (item_col, "VARCHAR(64)", "NOT NULL", desc),
    ]
schema["user_weapon_upgrades"] = [
    ("user_id", "BIGINT", "PK(part), FK users(id), CASCADE", "שחקן"),
    ("weapon_id", "VARCHAR(64)", "PK(part)", "נשק"),
    ("upgrade_level", "INTEGER", "NOT NULL", "רמת שדרוג"),
]
schema["user_map_plays"] = [
    ("user_id", "BIGINT", "PK(part), FK users(id), CASCADE", "שחקן"),
    ("map_id", "VARCHAR(64)", "PK(part)", "מפה"),
    ("plays", "INTEGER", "NOT NULL", "מספר משחקים"),
]
schema["user_weapon_kills"] = [
    ("user_id", "BIGINT", "PK(part), FK users(id), CASCADE", "שחקן"),
    ("weapon_id", "VARCHAR(64)", "PK(part)", "נשק"),
    ("kills", "INTEGER", "NOT NULL", "הריגות בנשק"),
]
for table_name, fields in schema.items():
    add_heading(doc, f"8.{list(schema).index(table_name)+1} טבלה: {table_name}", 2)
    add_table(doc, ["שם שדה", "סוג נתון", "הגדרות", "תיאור"], fields, widths=[1.45, 1.35, 1.8, 1.75], font_size=8.1)
add_heading(doc, "8.17 קשרים ואינדקסים", 2)
add_bullets(doc, [
    "כל טבלאות user_* מקושרות אל users באמצעות ON DELETE CASCADE.",
    "בקשות חברות, חברויות והזמנות כוללות שני קשרים אל users.",
    "room_invitations.room_code מקושר אל lobby_rooms.id.",
    "אינדקסים קיימים על username, token, צמדי משתמשים, בקשות נכנסות והזמנות נכנסות.",
    "מחיקת משתמש מוחקת קודים, קשרים חברתיים ואוספי התקדמות.",
])

add_heading(doc, "פרק 9 — עיצוב תפריטים", 1)
add_table(doc, ["סוג משתמש", "תפריטים זמינים", "הערות"], [
    ["אורח", "Loading → Auth → Login / Register / Verify Email", "אין גישה ללובי או ל־API מוגן ללא JWT."],
    ["משתמש רשום", "Main Menu, My Player, Settings, Rooms, Create Room, Match, Pause, Death Customizer", "זהו ה־Role המרכזי."],
    ["מנהל", "אותם מסכים של משתמש רשום", "בקוד אין תפריט ניהול נפרד; admin מעניק XP/ארנק גבוהים ומלאי מלא למשתמש seed."],
], widths=[1.4, 3.3, 1.65], font_size=8.8)
add_heading(doc, "9.1 תפריט ראשי", 2)
add_bullets(doc, [
    "Play — מעבר לחדרים ולבחירת משחק.",
    "My Player — פרופיל, תצוגת דמות, חנות, ציוד ושדרוגים.",
    "Settings — התאמת מקשי שליטה ואיפוס להגדרות ברירת מחדל.",
    "Daily Missions — כרטיסי משימה, התקדמות, Claim ו־Reroll.",
    "Sign Out — ניקוי הסשן וחזרה למסך האימות.",
])
add_heading(doc, "9.2 ניהול משתמשים והרשאות", 2)
add_table(doc, ["יכולת", "מצב", "מימוש שנמצא"], [
    ["Registration", "קיים", "POST /api/users/register יוצר משתמש, מבצע hashing ושולח קוד אימות."],
    ["Email Verification", "קיים", "קוד בן 6 ספרות, תוקף 15 דקות, טבלת email_verification_tokens."],
    ["Login", "קיים", "POST /api/users/login מאמת username/password וחוסם משתמש לא מאומת."],
    ["JWT Session", "קיים", "Bearer token נשמר ב־localStorage ונבדק ידנית בכל Controller מוגן."],
    ["Read Profile", "קיים", "GET /api/users/me."],
    ["Update Profile/Progress", "קיים חלקית", "PATCH /api/users/me/progress מעדכן התקדמות, ציוד וסטטיסטיקות; אינו עורך username/email."],
    ["Create User by Admin", "לא קיים", "אין endpoint או מסך מנהל ליצירת משתמש."],
    ["Update User by Admin", "לא קיים", "אין CRUD מנהלי."],
    ["Delete User", "לא קיים", "לא נמצא DELETE עבור משתמש."],
    ["List All Users", "לא קיים", "קיים רק חיפוש שחקנים חברתי לפי username."],
    ["Roles", "חלקי", "קיים boolean admin בלבד; אין טבלת Roles או enum תפקידים."],
    ["Permissions", "לא קיים", "לא נמצאה מערכת הרשאות פרטנית."],
    ["Admin Page", "לא קיים", "למנהל מוצגים אותם מסכים כמו למשתמש רגיל."],
], widths=[1.55, 1.15, 3.65], font_size=8.2)

add_heading(doc, "פרק 10 — עיצוב ויזואלי של מסכים", 1)
add_table(doc, ["מסך", "מרכיבים מרכזיים", "התנהגות"], [
    ["Loading", "לוגו, פס התקדמות, טעינת נכסים ובדיקת backend.", "ממתין למשך מתוזמן ולנגישות השירות."],
    ["Authentication", "בחירת Login/Register, טפסי פרטים, אימות קוד.", "מציג סטטוס והודעות שגיאה."],
    ["Main Menu", "מיתוג, דמות תלת־ממדית, רמה, XP, משימות ופקודות ניווט.", "מרכז הפעילות לאחר כניסה."],
    ["My Player", "CharacterPreview, לשוניות חנות, פריטי ציוד, מחירים ושדרוגים.", "Preview לפני רכישה/ציוד."],
    ["Settings", "טבלת פעולות ומקשים, מצב האזנה למקש.", "שמירה מקומית לכל משתמש."],
    ["Rooms", "רשימת חדרים, פרטי מפה ומצב, קוד חדר, Social Panel.", "רענון כל 3 שניות."],
    ["Create Room", "בחירת מפה, mode, score, time, max players, bots.", "ולידציה גם בלקוח וגם בשרת."],
    ["Match", "Canvas תלת־ממדי, כוונת, HUD, חיים, תחמושת, כסף, זמן וניקוד.", "עדכון פריים־אחר־פריים ועדכוני רשת."],
    ["Scoreboard", "קבוצות או כלל השחקנים, K/A/D/Score.", "מוצג בזמן החזקת Tab."],
    ["Pause/Death", "המשך, יציאה, התאמת דמות, החלפת נשק, kill cam.", "עצירה מקומית ואישור יציאה."],
], widths=[1.25, 3.25, 1.85], font_size=8.4)
add_heading(doc, "10.1 מובייל ו־Responsive", 2)
add_text(doc, "נמצאו viewport meta וכללי @media ברוחבים 720px, 760px ו־860px, המשנים grids, לובי, HUD, חנות והתאמת דמות. כלומר, קיימת התאמה ויזואלית חלקית למסכים צרים. עם זאת, מנגנון הקלט במשחק נשען על keyboard events, mouse events ו־Pointer Lock. לא נמצאו touchstart/touchmove/touchend, ג׳ויסטיק וירטואלי או כפתורי ירי למגע. לכן אין תמיכה מלאה במשחק במובייל.")
add_table(doc, ["בדיקה", "מצב"], [
    ["React Responsive", "קיים חלקית — הרכיבים משתמשים ב־CSS responsive."],
    ["Media Queries", "קיימים @media max-width: 720px, 760px ו־860px."],
    ["Responsive Layout", "קיים עבור חלק גדול ממסכי הלובי, החנות וה־HUD."],
    ["Touch Controls", "לא קיים בפרויקט הנוכחי."],
    ["Pointer Lock", "קיים ונדרש למשחק העיקרי, ולכן מתאים בעיקר למחשב."],
], widths=[2.0, 4.35], font_size=8.6)
add_heading(doc, "10.2 רשימת מסכים וצילומי מסך נדרשים", 2)
add_text(doc, "לא נמצאו קובצי Screenshot ייעודיים בפרויקט. קיימים לוגואים, מודלים, טקסטורות ונכסי UI, אך הם אינם תיעוד מצולם של המסכים. לפיכך יש לצלם את המסכים הבאים מתוך המערכת הפועלת:")
add_table(doc, ["מס׳", "מסך לצילום", "מה צריך להופיע בצילום"], [
    ["1", "Loading Screen", "לוגו DeadZone, פס טעינה ומצב backend."],
    ["2", "Authentication Landing", "אפשרויות Login ו־Create Account."],
    ["3", "Registration", "שם משתמש, דוא״ל, סיסמה ואימות סיסמה."],
    ["4", "Email Verification", "שדה קוד בן 6 ספרות והודעת מצב."],
    ["5", "Login", "שם משתמש, סיסמה והודעת שגיאה/הצלחה."],
    ["6", "Main Menu", "דמות, רמה, XP, ארנק, כפתורי ניווט ומשימות יומיות."],
    ["7", "My Player / Store", "תצוגת דמות, לשוניות תלבושות, נשקים, סקינים ואביזרים."],
    ["8", "Purchase Confirmation", "חלון אישור רכישה, מחיר ופריט."],
    ["9", "Settings", "מיפוי מקשים וכפתור איפוס."],
    ["10", "Rooms", "רשימת חדרים, מפה, מצב משחק, תפוסה וקוד חדר."],
    ["11", "Create Room", "בחירת מפה, mode, score, time, max players ו־bots."],
    ["12", "Social — Friends", "חיפוש שחקנים ורשימת חברים."],
    ["13", "Social — Requests", "בקשות נכנסות/יוצאות וכפתורי Accept/Decline."],
    ["14", "Social — Room Invites", "הזמנות לחדר וכפתורי קבלה/דחייה."],
    ["15", "Match HUD", "זירת 3D, כוונת, חיים, תחמושת, רימונים, כסף, זמן וניקוד."],
    ["16", "Scoreboard", "K/A/D/Score לכל שחקן או קבוצה."],
    ["17", "Pause Menu", "Continue, התאמת דמות ו־Exit Match."],
    ["18", "Kill Cam / Death Screen", "שם הפוגע, זמן respawn ונתוני השחקן."],
    ["19", "Respawn Customizer", "שינוי תלבושת, אביזרים ונשק בזמן מוות."],
    ["20", "Match Result", "Victory/Draw/Match Complete ותגמול ניצחון."],
    ["21", "Exit Confirmation", "Stay ו־Exit."],
], widths=[0.55, 1.75, 4.05], font_size=8.0)

add_heading(doc, "פרק 11 — תרשים זרימה בין מסכים", 1)
add_picture(doc, flow_img, Inches(6.35), "תרשים זרימת מסכים של DeadZone")
add_text(doc, "תרשים 4 — זרימת מסכים עיקרית", bold=True, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
add_code(doc, """
flowchart LR
    Loading --> Auth
    Auth --> Login
    Auth --> Register
    Register --> VerifyEmail
    Login --> MainMenu
    VerifyEmail --> Login
    MainMenu --> MyPlayer
    MainMenu --> Settings
    MainMenu --> Rooms
    Rooms --> CreateRoom
    Rooms --> Match
    Match --> Scoreboard
    Match --> Pause
    Match --> DeathScreen
    Pause --> Rooms
    DeathScreen --> Match
""")

add_heading(doc, "פרק 12 — שאילתות ודוחות", 1)
add_heading(doc, "12.1 פעולות Repository מרכזיות", 2)
add_table(doc, ["מטרה", "פעולה/שאילתה לוגית", "תוצאה"], [
    ["כניסה", "SELECT user WHERE username = ?", "משתמש יחיד או 401."],
    ["חיפוש שחקנים", "username contains query, excluding current user", "רשימת UserSummary."],
    ["בקשות נכנסות", "recipient_id = ? AND status = 'PENDING'", "בקשות ממתינות."],
    ["חברים", "first_user_id = ? OR second_user_id = ?", "קשרי חברות."],
    ["הזמנות פעילות", "recipient_id = ? AND status = 'PENDING' AND expires_at > now", "הזמנות חדר."],
    ["חדרים פתוחים", "SELECT all rooms; filter players < max_players; sort by name", "רשימת חדרים."],
    ["ניקוי חדרים", "non-permanent AND players=0 AND last_activity < cutoff", "מחיקת חדרים זמניים."],
    ["סטטיסטיקת משתמש", "users + element collections by user_id", "פרופיל מלא."],
], widths=[1.25, 3.6, 1.5], font_size=8.5)
add_heading(doc, "12.2 שאילתות SQL לדוגמה", 2)
add_code(doc, """
-- חדרים פתוחים
SELECT * FROM lobby_rooms
WHERE players < max_players
ORDER BY name;

-- חברים של משתמש
SELECT * FROM friendships
WHERE first_user_id = :userId OR second_user_id = :userId;

-- הזמנות פעילות
SELECT * FROM room_invitations
WHERE recipient_id = :userId
  AND status = 'PENDING'
  AND expires_at > CURRENT_TIMESTAMP;

-- סיכום התקדמות משתמש
SELECT username, total_kills, total_assists, total_deaths, wallet, xp
FROM users
WHERE id = :userId;
""")
add_heading(doc, "12.3 דוחות ניהוליים מוצעים", 2)
add_callout(doc, "סטטוס מימוש", "Dashboard ניהולי, דוחות, גרפים, Charts ועמודי Analytics אינם קיימים בפרויקט הנוכחי. הטבלה הבאה היא המלצה להוספה בעתיד בלבד.")
add_table(doc, ["דוח", "עמודות דוגמה", "מקור"], [
    ["דירוג שחקנים", "username, kills, assists, deaths, K/D, xp", "users"],
    ["שימוש במפות", "map_id, SUM(plays), COUNT(users)", "user_map_plays"],
    ["פופולריות נשקים", "weapon_id, SUM(kills), AVG(upgrade_level)", "user_weapon_kills + user_weapon_upgrades"],
    ["חדרים פעילים", "room id, mode, map, players/max, activity", "lobby_rooms"],
    ["משימות", "mission_id, מספר תביעות", "user_claimed_missions"],
], widths=[1.35, 3.35, 1.65], font_size=8.5)
add_heading(doc, "12.4 מצב רכיבי ניהול ודיווח", 2)
add_table(doc, ["רכיב", "האם קיים?", "הערה"], [
    ["Dashboard ניהולי", "לא", "התפריט הראשי הוא Dashboard שחקן, לא מערכת ניהול."],
    ["Statistics לשחקן", "חלקי", "נשמרים kills, assists, deaths, XP, wallet ונתוני משימות; אין מסך אנליטי מלא."],
    ["Admin Pages", "לא", "לא נמצא Route או Component ייעודי למנהל."],
    ["Reports", "לא", "אין endpoints או מסכי דוחות."],
    ["Charts", "לא", "אין ספריית charts ואין רכיבי גרפים."],
    ["Export CSV/PDF", "לא", "לא קיים בפרויקט הנוכחי."],
], widths=[1.7, 1.15, 3.5], font_size=8.5)

add_heading(doc, "פרק 13 — בדיקות QA", 1)
qa = [
    ("QA-01", "רישום תקין", "שם/דוא״ל/סיסמה חדשים", "נוצר משתמש לא מאומת וקוד", "עבר בבדיקת UserService"),
    ("QA-02", "סיסמה מוצפנת", "secret", "הערך במסד שונה ומתאים בבדיקה", "עבר"),
    ("QA-03", "אימות דוא״ל", "קוד תקין לפני תפוגה", "email_verified=true", "לא קיימת בדיקה אוטומטית ייעודית"),
    ("QA-04", "חסימת כניסה לא מאומתת", "משתמש לא מאומת", "תגובה מתאימה ללא JWT", "נדרש להשלים בדיקת integration"),
    ("QA-05", "JWT פגום", "Bearer invalid", "401", "נדרש להשלים"),
    ("QA-06", "יצירת חדר objective", "capture-flag, score 9, 15 min", "הערכים נשמרים", "עבר"),
    ("QA-07", "חדר מלא", "join כאשר players=max", "409 Conflict", "נדרש להשלים"),
    ("QA-08", "סניטציית שחקן", "שם ארוך, red, smg", "שם נחתך ל־16 ושדות נשמרים", "עבר"),
    ("QA-09", "עזיבת חדר", "השחקן האחרון יוצא", "GameRoom נמחק מהזיכרון", "עבר"),
    ("QA-10", "פגיעה באויב", "damage=100", "מוות, kill ו־score מתעדכנים", "עבר"),
    ("QA-11", "Friendly Fire", "שני שחקנים באותה קבוצה", "אין נזק", "מכוסה חלקית בבדיקה"),
    ("QA-12", "שמירת התקדמות", "מלאי, XP, stats, upgrades", "כל השדות נשמרים", "עבר"),
    ("QA-13", "מניעת רכישה ללא תשלום", "מלאי חדש בלי הפחתת ארנק", "400", "עבר"),
    ("QA-14", "חיפוש משתמש קצר", "תו אחד", "אין בקשה לשרת/רשימה ריקה", "נדרש UI test"),
    ("QA-15", "בקשת חברות כפולה", "אותו sender/recipient", "דחייה באמצעות unique/business rule", "נדרש integration"),
    ("QA-16", "הזמנה שפגה", "expires_at בעבר", "לא מופיעה כפעילה/לא מתקבלת", "נדרש integration"),
    ("QA-17", "Reroll חינם ראשון", "count=0", "עלות 0", "נדרש unit test client"),
    ("QA-18", "פתיחת נשק לפי Level", "רמה נמוכה", "כפתור נעול", "נדרש UI test"),
    ("QA-19", "Build צד לקוח", "npm/vite build", "Build תקין", "עבר; קיימות אזהרות"),
    ("QA-20", "כל בדיקות שרת", "mvn test", "11/11 ללא כשל", "עבר ב־25.6.2026"),
]
add_table(doc, ["מס׳", "שם בדיקה", "קלט/תיאור", "תוצאה צפויה", "סטטוס"], qa,
          widths=[0.55, 1.2, 1.85, 1.75, 1.0], font_size=7.7)
add_heading(doc, "13.1 תוצאות הרצה בפועל", 2)
add_bullets(doc, [
    "Maven: 11 בדיקות, 0 כשלים, 0 שגיאות, 0 דילוגים — BUILD SUCCESS.",
    "Vite: בניית production הושלמה בהצלחה.",
    "אזהרה: bundle JavaScript כ־958KB לפני gzip; מומלץ code splitting.",
    "אזהרה: שלושה קובצי רקע המוזכרים ב־CSS לא נמצאו בזמן build: login_card_bg.png, grid_pattern.png, metal_background.png.",
])

add_heading(doc, "פרק 14 — טכנולוגיות", 1)
add_table(doc, ["שכבה", "טכנולוגיה", "שימוש בפרויקט"], [
    ["Frontend", "React 19.2", "רכיבי מסך, state ו־hooks."],
    ["Build", "Vite 8", "שרת פיתוח ובניית production."],
    ["3D", "Three.js 0.184", "Scene, Camera, WebGLRenderer, meshes, fog, lighting."],
    ["React 3D", "React Three Fiber / Drei", "תצוגת דמות ומודלים תלת־ממדיים ברכיבים."],
    ["Physics", "React Three Rapier", "תלות פיזיקלית; עיקר התנגשות המשחק ממומש במערכות קוד ייעודיות."],
    ["Backend", "Java 21 + Spring Boot 4.1", "REST, DI, services, scheduling ותצורה."],
    ["API", "Spring Web MVC", "Controllers ו־JSON."],
    ["Realtime", "Spring WebSocket", "חיבור /game וסנכרון חדרים."],
    ["ORM", "Spring Data JPA / Hibernate 7", "Entities, repositories ו־element collections."],
    ["Database", "PostgreSQL / Supabase", "מסד הייצור."],
    ["Migrations", "Flyway", "יצירת סכמה, FK ואינדקסים."],
    ["Offline/Test DB", "H2", "הרצה מקומית ובדיקות."],
    ["Authentication", "JWT HS256 מותאם", "סשן Bearer ל־30 יום."],
    ["Password", "Java SHA-256 + salt", "שמירת סיסמאות ללא טקסט גלוי."],
    ["Email", "Spring Mail / SMTP", "שליחת קוד אימות."],
    ["Serialization", "Jackson", "JSON ב־REST וב־WebSocket."],
    ["Deployment", "Dockerfiles + env vars", "קיימים קובצי Docker לשני הצדדים."],
    ["Storage", "localStorage", "JWT ומיפוי מקשים; נתוני התקדמות במסד."],
], widths=[1.15, 2.05, 3.15], font_size=8.3)
add_heading(doc, "14.1 כל חבילות צד הלקוח — package.json", 2)
add_table(doc, ["סוג", "חבילה", "גרסה", "תפקיד"], [
    ["Dependency", "@react-three/drei", "^10.7.7", "Helpers ורכיבי עזר ל־React Three Fiber."],
    ["Dependency", "@react-three/fiber", "^9.6.1", "Renderer של Three.js בתוך React."],
    ["Dependency", "@react-three/rapier", "^2.2.0", "אינטגרציית Rapier physics ל־React Three Fiber."],
    ["Dependency", "react", "^19.2.6", "ספריית ממשק משתמש."],
    ["Dependency", "react-dom", "^19.2.6", "רינדור React ל־DOM."],
    ["Dependency", "three", "^0.184.0", "מנוע גרפי תלת־ממדי WebGL."],
    ["Dev", "@eslint/js", "^10.0.1", "כללי ESLint ל־JavaScript."],
    ["Dev", "@types/react", "^19.2.14", "טיפוסים וכלי פיתוח עבור React."],
    ["Dev", "@types/react-dom", "^19.2.3", "טיפוסים וכלי פיתוח עבור React DOM."],
    ["Dev", "@vitejs/plugin-react", "^6.0.1", "שילוב React ב־Vite."],
    ["Dev", "eslint", "^10.3.0", "בדיקות lint."],
    ["Dev", "eslint-plugin-react-hooks", "^7.1.1", "כללי Hooks."],
    ["Dev", "eslint-plugin-react-refresh", "^0.5.2", "כללי Fast Refresh."],
    ["Dev", "globals", "^17.6.0", "הגדרות globals ל־ESLint."],
    ["Dev", "vite", "^8.0.12", "שרת פיתוח ובניית production."],
], widths=[0.8, 2.1, 1.0, 2.45], font_size=7.9)
add_heading(doc, "14.2 כל תלויות צד השרת — pom.xml", 2)
add_table(doc, ["Dependency", "Scope", "תפקיד"], [
    ["spring-boot-starter-webmvc", "compile", "REST Controllers ו־Spring MVC."],
    ["spring-boot-starter-websocket", "compile", "WebSocket endpoint ומשחק בזמן אמת."],
    ["spring-boot-starter-webmvc-test", "test", "בדיקות Spring MVC."],
    ["spring-boot-starter-websocket-test", "test", "בדיקות WebSocket."],
    ["spring-boot-starter-data-jpa", "compile", "Spring Data JPA ו־Hibernate ORM."],
    ["spring-boot-starter-mail", "compile", "SMTP ואימות דוא״ל."],
    ["postgresql", "runtime", "JDBC driver ל־PostgreSQL/Supabase."],
    ["spring-boot-starter-flyway", "compile", "אינטגרציית migration."],
    ["flyway-database-postgresql", "compile", "תמיכת Flyway ב־PostgreSQL."],
    ["h2", "runtime", "מסד offline ובדיקות."],
    ["jackson-databind", "compile", "JSON serialization/deserialization."],
    ["spring-boot-maven-plugin", "build plugin", "אריזת והרצת יישום Spring Boot."],
], widths=[2.65, 1.15, 2.55], font_size=8.2)
add_heading(doc, "14.3 Deployment והרצת המערכת", 2)
add_table(doc, ["נושא", "ממצא מבוסס קוד"], [
    ["הרצת Client בפיתוח", "מתיקיית deadzone-client: npm ci או npm install, ולאחר מכן npm run dev. קיימת גם פקודת npm run dev:lan."],
    ["Client Port", "Vite אינו מקבע port בקוד; ברירת המחדל של Vite היא 5173 אם הפורט פנוי. APP_BASE_URL גם מוגדר כברירת מחדל ל־http://localhost:5173."],
    ["הרצת Server", "מתיקיית shooter-server: .\\mvnw.cmd spring-boot:run ב־Windows או ./mvnw spring-boot:run."],
    ["Server Port", "server.port=${PORT:8080}; ברירת המחדל היא 8080."],
    ["Client API Origin", "VITE_API_ORIGIN; אם חסר, הלקוח פונה לאותו hostname בפורט 8080."],
    ["Database Production", "PostgreSQL באמצעות DB_URL, DB_USERNAME ו־DB_PASSWORD; קובצי הסביבה מצביעים על Supabase."],
    ["Database Offline", "application-offline.properties משתמש ב־H2 file במצב תאימות PostgreSQL, ddl-auto=update וללא Flyway."],
    ["Database Tests", "H2 in-memory, create-drop וללא Flyway."],
    ["Docker Client", "Multi-stage: Node 22 Alpine build ולאחריו nginx 1.27 Alpine; container חושף port 80."],
    ["Docker Server", "Multi-stage Maven/Temurin 21; JRE non-root user; container חושף port 8080."],
    ["Docker Compose", "לא נמצא docker-compose.yml, compose.yml או קובץ Compose אחר."],
    ["Production Profile", "לא נמצא application-prod.properties/yml ולא קיים Spring profile ייעודי בשם production. application.properties הראשי מוגדר עבור PostgreSQL ו־Flyway."],
    ["application.yml", "לא נמצא; התצורה נמצאת בקובצי .properties."],
], widths=[1.75, 4.6], font_size=8.2)
add_heading(doc, "14.4 משתני סביבה", 2)
add_table(doc, ["משתנה", "שימוש"], [
    ["VITE_API_ORIGIN", "כתובת בסיס השרת בזמן build/run של הלקוח."],
    ["PORT", "פורט שרת; ברירת מחדל 8080."],
    ["DB_URL", "JDBC URL של PostgreSQL."],
    ["DB_USERNAME", "משתמש מסד נתונים."],
    ["DB_PASSWORD", "סיסמת מסד נתונים."],
    ["DB_POOL_SIZE", "גודל מרבי של Hikari pool; ברירת מחדל 8."],
    ["DB_POOL_MIN_IDLE", "מספר חיבורים idle מינימלי; ברירת מחדל 1."],
    ["DB_CONNECTION_TIMEOUT", "connection timeout; ברירת מחדל 15000ms."],
    ["JWT_SECRET", "מפתח חתימת JWT."],
    ["APP_BASE_URL", "כתובת הלקוח; ברירת מחדל localhost:5173."],
    ["MAIL_FROM", "כתובת From; יכולה לרשת SMTP_USERNAME."],
    ["SMTP_HOST / SMTP_PORT", "שרת SMTP; ברירות מחדל smtp.gmail.com:587."],
    ["SMTP_USERNAME / SMTP_PASSWORD", "פרטי התחברות SMTP."],
    ["SMTP_AUTH", "הפעלת SMTP authentication."],
    ["SMTP_STARTTLS / SMTP_STARTTLS_REQUIRED", "הפעלת ודרישת STARTTLS."],
    ["SMTP_CONNECTION_TIMEOUT", "ברירת מחדל 10000ms."],
    ["SMTP_TIMEOUT", "ברירת מחדל 10000ms."],
    ["SMTP_WRITE_TIMEOUT", "ברירת מחדל 10000ms."],
    ["EMAIL_DELIVERY_REQUIRED", "אם false, קוד האימות נכתב ללוג במקום לחייב SMTP."],
    ["SUPABASE_URL", "מוגדר בתצורה; לא נמצא שימוש ישיר בשירותי Supabase API בקוד Java."],
    ["SUPABASE_PUBLISHABLE_KEY", "מוגדר בתצורה; לא נמצא שימוש ישיר בקוד."],
    ["SUPABASE_SECRET_KEY", "מוגדר בתצורה; לא נמצא שימוש ישיר בקוד."],
    ["SUPABASE_JWKS_URL", "מוגדר בתצורה; JwtService הנוכחי אינו משתמש ב־JWKS."],
], widths=[2.35, 4.0], font_size=7.9)

add_heading(doc, "פרק 15 — קריטריונים להצלחת המערכת", 1)
add_table(doc, ["מדד", "יעד קבלה", "אופן בדיקה"], [
    ["זמן תגובת REST", "95% מהבקשות עד 500ms בסביבה רגילה", "מדידת API תחת עומס מוגדר."],
    ["זמן עד לובי", "עד 5 שניות לאחר טעינת נכסים ברשת תקינה", "Performance trace."],
    ["יציבות רינדור", "45 FPS לפחות במחשב יעד; יעד 60 FPS", "מדידת FPS בשלוש מפות."],
    ["תקינות קרב", "אין friendly fire במצבי קבוצה; ניקוד ומוות עקביים", "בדיקות משחק מרובות לקוחות."],
    ["אבטחה", "100% endpoints מוגנים דוחים JWT חסר/פגום", "Integration tests."],
    ["רכישות", "אין אפשרות להוסיף פריט ללא תשלום שרת תקין", "בדיקות מניפולציית payload."],
    ["שימושיות", "80% ממשתמשי ניסוי נכנסים למשחק ללא הדרכה", "בדיקת משתמשים."],
    ["אמינות", "0 אובדן התקדמות לאחר יציאה/כניסה", "בדיקות התמדה."],
    ["התאמה לדרישות", "כל RF-01–RF-23 מאומתות או מסומנות מחוץ להיקף", "Traceability matrix."],
    ["איכות קוד", "Build נקי מכשלים וכל בדיקות השרת עוברות", "CI אוטומטי."],
], widths=[1.4, 3.05, 1.9], font_size=8.5)

add_heading(doc, "פרק 16 — לוח זמנים", 1)
add_callout(doc, "בסיס עובדתי", "לוח הזמנים מבוסס על היסטוריית Git המקומית של המאגר origin=https://github.com/mooshnick/DeadZone.git. ההיסטוריה המתועדת כוללת commits בין 15 ל־24 ביוני 2026. לא קיימים timesheets, שעות עבודה או מסמך Sprint, ולכן לא ניתן לקבוע שעות בפועל.")
add_table(doc, ["תאריך/טווח", "שלב", "ראיות מהיסטוריית Git", "תוצרים שניתן לאמת"], [
    ["15.06.2026", "אפיון ותשתית ראשונית", "Initial commit; Server and Client; העלאת שרת ודאטה בייס.", "מבנה React/Vite, Spring Boot, WebSocket ראשוני, User ו־Repository."],
    ["16.06.2026", "Frontend + Backend", "שינויים ב־App, CSS, package.json, Controllers, Models ו־WebSocket.", "אינטגרציה ראשונית בין ממשק, API ושרת."],
    ["17.06.2026", "Multiplayer ומנוע משחק", "נוספו GameWorld, מערכות Combat/Collision/Bot/Grenade/Powerup, ArenaLayouts ושירותי חדר בזמן אמת.", "מנוע המשחק המרכזי, שחקנים, בוטים ותקשורת חדר."],
    ["18–20.06.2026", "לא מתועד", "לא נמצאו commits בתאריכים אלה.", "לא ניתן לקבוע פעילות מתוך המאגר."],
    ["21.06.2026", "Backend, אבטחה ו־QA", "נוספו DTOs, PasswordService, UserService ובדיקות User/GameRules/Room.", "hashing, שכבת שירות, תגובות API ובדיקות אוטומטיות."],
    ["22.06.2026", "לא מתועד", "לא נמצאו commits בתאריך זה.", "לא ניתן לקבוע פעילות מתוך המאגר."],
    ["23.06.2026", "Design, Assets, Rooms ו־Game Modes", "נוספו לוגואים, מודל העיר, ערכות 3D, LobbyRoom, API חדרים ועדכוני HUD/GameWorld.", "מיתוג, מפות ונכסים, יצירת חדרים, מצבי משחק ושיפור הממשק."],
    ["24.06.2026", "Database, Email, Social, Deployment וייצוב Multiplayer", "commits: מייל עובד, Dockerfiles, שרת עובד, חברים, דאטה בייס חיצוני והחלפת MYSQL.", "אימות דוא״ל, JWT/realtime client, מערכת חברים, Docker, PostgreSQL/Flyway ו־H2 offline."],
    ["25.06.2026", "QA ותיעוד", "בדיקות שבוצעו בסביבת העבודה: Maven ו־Vite build.", "11 בדיקות שרת עברו; build צד לקוח הצליח; הוכן דוח הפרויקט."],
], widths=[1.0, 1.45, 2.45, 1.45], font_size=7.7)
add_heading(doc, "16.1 חלוקה לפי תחומי העבודה שנדרשו", 2)
add_table(doc, ["תחום", "תקופה מתועדת", "היקף שניכר בקוד"], [
    ["אפיון", "15.06 ואילך", "מבנה מערכת, משתמש, חדר, שחקן ומסרים; מסמך אפיון נפרד לא נמצא."],
    ["Design", "23–24.06", "CSS רחב, לוגואים, HUD, חנות, מסכי לובי ונכסי 3D."],
    ["Frontend", "15–24.06", "React SPA, ניתוב ידני, מסכים, state, API, HUD וחנות."],
    ["Backend", "15–24.06", "Controllers, Services, DTOs, repositories, email ו־business validation."],
    ["Database", "15, 21, 23–24.06", "H2 ראשוני, JPA, PostgreSQL, Flyway, סכמה מנורמלת."],
    ["Multiplayer", "17 ו־24.06", "WebSocket, rooms, state broadcast, RealtimeClient ועדכון תנועה 50ms."],
    ["QA", "21 ו־25.06", "בדיקות יחידה/אינטגרציה והרצות build; בדיקות UI ועומס לא נמצאו."],
    ["Deployment", "24.06", "Dockerfiles ו־ENV; Docker Compose ו־production profile ייעודי אינם קיימים."],
], widths=[1.25, 1.45, 3.65], font_size=8.3)

add_heading(doc, "פרק 17 — נספחים, תרשימים וביבליוגרפיה", 1)
add_heading(doc, "17.1 רשימת תרשימים", 2)
add_numbered(doc, [
    "תרשים 1 — ERD מרכזי של DeadZone.",
    "תרשים 2 — סקירת תרחישי שימוש.",
    "תרשים 3 — ארכיטקטורת המערכת.",
    "תרשים 4 — זרימת מסכים עיקרית.",
])
add_heading(doc, "17.2 רשימת טבלאות", 2)
add_bullets(doc, [
    "טבלת תקציר טכנולוגי ומאפייני מערכת.",
    "השוואת מערכות דומות.",
    "דרישות פונקציונליות ולא פונקציונליות.",
    "ישויות, תרחישי שימוש, תפריטים ומסכים.",
    "טבלאות סכמת מסד הנתונים.",
    "שאילתות ודוחות, בדיקות QA, טכנולוגיות, מדדי הצלחה ולוח זמנים.",
])
add_heading(doc, "17.3 ביבליוגרפיה", 2)
add_bullets(doc, [
    "קוד המקור של DeadZone: deadzone-client ו־shooter-server.",
    "מאגר GitHub של הפרויקט — https://github.com/mooshnick/DeadZone.git",
    "Spring Boot Reference Documentation — https://docs.spring.io/spring-boot/",
    "React Documentation — https://react.dev/",
    "Three.js Documentation — https://threejs.org/docs/",
    "PostgreSQL Documentation — https://www.postgresql.org/docs/",
    "VALORANT official site — https://playvalorant.com/",
    "Counter-Strike 2 official site — https://www.counter-strike.net/cs2",
    "Krunker official site — https://krunker.io/",
    "Fortnite official site — https://www.fortnite.com/",
])
add_heading(doc, "17.4 נספחי קוד חשובים", 2)
add_code(doc, """
// REST routes
POST   /api/users/register
POST   /api/users/login
POST   /api/users/verify-email
GET    /api/users/me
PATCH  /api/users/me/progress
GET    /api/rooms
POST   /api/rooms
POST   /api/rooms/{code}/join
POST   /api/rooms/{code}/leave
GET    /api/social
GET    /api/social/users?username=...
POST   /api/social/friend-requests
POST   /api/social/room-invites

// Realtime endpoint
WS /game: JOIN, MOVE, SHOOT, HIT -> STATE, SHOT, ERROR
""")
add_heading(doc, "17.5 Credits ונכסי הפרויקט", 2)
add_callout(doc, "רישוי וייחוס", "לא נמצאו קובצי LICENSE, attribution או credits בתוך תיקיות הנכסים. לכן לא ניתן לקבוע את יוצרי המודלים, הטקסטורות, הלוגואים או הצליל מתוך הקוד בלבד. יש להשלים מקור ורישיון לפני פרסום.")
add_table(doc, ["קטגוריה", "כמות/פרטים שנמצאו", "מיקום והערות"], [
    ["כלל קובצי הנכסים", "659 קבצים", "deadzone-client/public ו־deadzone-client/src/assets."],
    ["PNG", "128", "לוגואים, hero, טקסטורות castle_kit, kit_2 ו־kit_3."],
    ["SVG / Icons", "4", "favicon.svg, icons.svg, react.svg, vite.svg."],
    ["GLB Models", "132", "39 ב־castle_kit, 91 ב־kit_3 ושני מודלי apocalyptic_city."],
    ["FBX Models", "130", "39 ב־castle_kit ו־91 ב־kit_3."],
    ["OBJ + MTL", "130 OBJ ו־130 MTL", "חלופות מודל וחומר עבור castle_kit ו־kit_3."],
    ["Textures", "128 PNG בסך הכול", "117 קובצי kit_2, קובצי colormap ו־variation בערכות האחרות, וכן לוגואים/hero."],
    ["Sound", "קובץ MP3 אחד", "public/sound/shotSound_1.mp3."],
    ["Draco", "3 קובצי JS וקובץ WASM אחד", "decoder/encoder לדחיסת מודלי glTF."],
    ["Fonts", "אין קובצי font", "CSS משתמש ב־Inter, system-ui, Segoe UI, Impact, Arial Black, Consolas ו־Courier New דרך גופני מערכת."],
    ["Generated Tool", "gltfjsx", "ApocalypticCityModel.jsx כולל הערת Auto-generated by github.com/pmndrs/gltfjsx."],
    ["AI Tools", "לא נמצא בקוד", "אין תיעוד לכלי AI ששימשו. אין לייחס כלי AI ללא מידע חיצוני מהצוות."],
], widths=[1.4, 1.55, 3.4], font_size=8.0)
add_heading(doc, "17.6 משפחות נכסים", 2)
add_bullets(doc, [
    "deadZone_Logo.png ו־Shadow_Logo.png — מיתוג המערכת.",
    "favicon.svg ו־icons.svg — favicon וספריית סמלים.",
    "src/assets/hero.png — נכס תמונת hero; react.svg ו־vite.svg — נכסי תבנית.",
    "public/model — apocalyptic_city.glb ו־apocalyptic_city-transformed.glb.",
    "public/castle_kit — 39 מודלים בכל אחד מהפורמטים GLB/FBX/OBJ, קובצי MTL וטקסטורות.",
    "public/kit_2 — 117 טקסטורות PNG של רצפות, קירות, גגות, דלתות וחלונות.",
    "public/kit_3 — 91 מודלים בכל אחד מהפורמטים GLB/FBX/OBJ, MTL וטקסטורות; משפחות cemetery, crypt, fence, vegetation ו־props.",
    "public/sound — צליל ירייה יחיד.",
    "public/draco — decoder, encoder ו־WASM עבור מודלים דחוסים.",
])

add_heading(doc, "פרק 18 — English Abstract", 1)
add_text(doc, "DeadZone is a browser-based multiplayer first-person shooter that combines a React and Three.js 3D client with a Java Spring Boot backend. The system supports account registration, email verification, JWT-based sessions, persistent player progression, a virtual store, daily missions, social features, room creation, bots, nine maps, six weapons, and five game modes. REST APIs handle users, rooms, social interactions, and progression, while a WebSocket endpoint synchronizes active matches. Production data is stored in PostgreSQL through Spring Data JPA and Flyway, with H2 available for offline development and automated testing. The project's main strength is the integration of realtime gameplay with persistent identity, progression, customization, and social play in a single browser-accessible application.", rtl=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading(doc, "רשימת פרטים להשלמה ידנית", 1)
add_table(doc, ["מס׳", "פרט חסר", "מה יש להשלים"], [
    ["1", "שם הקורס", "שם מלא ומספר קורס."],
    ["2", "שמות מגישים", "שמות מלאים, מספרי זהות/סטודנט לפי הנדרש."],
    ["3", "מנחה/מרצה", "שם המרצה והמוסד."],
    ["4", "תאריך הגשה", "לאשר או להחליף את 25 ביוני 2026."],
    ["5", "דוח הדוגמה", "לא נמצא קובץ הדוח 'עת לעשות' בתיקיית העבודה; ניתן להתאים עיצוב אם יצורף."],
    ["6", "שעות עבודה בפועל", "Git מספק תאריכים ושינויים, אך לא שעות עבודה או חלוקה אישית בין המגישים."],
    ["7", "צילומי מסך", "לא נמצאו screenshots בפרויקט. יש לצלם את 21 המסכים המפורטים בפרק 10.2."],
    ["8", "כתובות פריסה בפועל", "לא נמצאו URL פעילים של Client/Server או שם ספק deployment בקוד."],
    ["9", "מדדי ביצועים", "יש להריץ מדידות FPS, latency, עומס וזמני תגובה; אין תוצאות בקוד."],
    ["10", "מדיניות גיבוי ושחזור", "לא נמצא מסמך או קוד המתאר backup/restore של PostgreSQL/Supabase."],
    ["11", "רישיונות וקרדיטים לנכסים", "לא נמצאו קובצי LICENSE/attribution לערכות 3D, טקסטורות, לוגואים והצליל."],
    ["12", "שימוש בכלי AI", "לא נמצא תיעוד בקוד. יש להשלים רק אם הצוות אכן השתמש בכלים כאלה."],
], widths=[0.55, 1.55, 4.25], font_size=8.3)
add_heading(doc, "המלצות לשיפור הפרויקט", 1)
add_bullets(doc, [
    "להחליף את מנגנון SHA-256 המותאם ב־bcrypt או Argon2 באמצעות ספריית אבטחה מקובלת.",
    "להוסיף Spring Security, הרשאות מבוססות Role, הגנת endpoints ותהליך ביטול/רענון JWT.",
    "לאמת JWT גם בחיבור WebSocket, להגביל Allowed Origins ולא לסמוך על playerId שמגיע מהלקוח.",
    "להגביל CORS לכתובות הפריסה בפועל ולהוסיף Rate Limiting ל־login, registration, verification ו־WebSocket.",
    "להוסיף בדיקות Integration ל־Controllers, JWT, אימות דוא״ל, תפוגת הזמנות, חדר מלא ובקשות חברתיות כפולות.",
    "להוסיף בדיקות UI/E2E, בדיקות עומס ומדידות FPS, latency וזמני תגובת API.",
    "להוסיף interpolation, snapshot buffering, reconnect/backoff ומדידת RTT למשחק מרובה־משתתפים.",
    "לשקול tick שרת authoritative ומנגנון משותף לחדרים כאשר מריצים יותר ממופע שרת אחד.",
    "לפצל את useDeadzoneController, Lobby.jsx ו־bundle הלקוח למודולים ו־chunks קטנים יותר.",
    "להוסיף Docker Compose או תיעוד deployment מלא, health checks, monitoring ומדיניות backup/restore.",
    "להוסיף ממשק Admin ודוחות רק אם הם דרישת מערכת; כיום הם אינם קיימים.",
    "להוסיף touch controls רק אם נדרשת תמיכה מלאה במובייל; כרגע התמיכה היא ויזואלית חלקית בלבד.",
    "לצרף קובצי רישיון ו־attribution לכל מודל, טקסטורה, סמל, לוגו וצליל.",
    "להוסיף את 21 צילומי המסך המפורטים בדוח לאחר הרצת המערכת בסביבת ההגשה.",
])
add_callout(doc, "סיכום מסירה", "המסמך משקף את מצב הקוד שנבדק. כל יכולת שאינה קיימת בפועל סומנה כהצעה, מגבלה או השלמה ידנית.")

# Document properties
doc.core_properties.title = "דוח פרויקט DeadZone"
doc.core_properties.subject = "דוח סיכום ותכנון פרויקט תוכנה"
doc.core_properties.author = "לא זוהה בקוד — להשלמה ידנית"
doc.core_properties.keywords = "DeadZone, React, Three.js, Spring Boot, WebSocket, PostgreSQL"

doc.save(OUT)
print(OUT)
